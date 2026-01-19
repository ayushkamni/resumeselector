"""
Text Extraction Module
Handles extraction of text from various file formats (PDF, DOCX)
"""

import PyPDF2
from docx import Document
import io
import streamlit as st

def extract_text_from_file(uploaded_file) -> str:
    """
    Extract text from uploaded file (PDF or DOCX)

    Args:
        uploaded_file: Streamlit uploaded file object

    Returns:
        str: Extracted text content
    """
    try:
        # Get file extension
        file_name = uploaded_file.name.lower()

        if file_name.endswith('.pdf'):
            return extract_text_from_pdf(uploaded_file)
        elif file_name.endswith('.docx'):
            return extract_text_from_docx(uploaded_file)
        else:
            raise ValueError("Unsupported file format. Please upload PDF or DOCX files.")

    except Exception as e:
        st.error(f"Error extracting text: {str(e)}")
        return ""

def extract_text_from_pdf(uploaded_file) -> str:
    """
    Extract text from PDF file

    Args:
        uploaded_file: PDF file object

    Returns:
        str: Extracted text
    """
    try:
        # Create PDF reader object
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))

        # Extract text from all pages
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"

        return text.strip()

    except Exception as e:
        raise Exception(f"PDF extraction failed: {str(e)}")

def extract_text_from_docx(uploaded_file) -> str:
    """
    Extract text from DOCX file

    Args:
        uploaded_file: DOCX file object

    Returns:
        str: Extracted text
    """
    try:
        # Load the document
        doc = Document(io.BytesIO(uploaded_file.read()))

        # Extract text from all paragraphs
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"

        return text.strip()

    except Exception as e:
        raise Exception(f"DOCX extraction failed: {str(e)}")

def get_file_info(uploaded_file) -> dict:
    """
    Get basic information about the uploaded file

    Args:
        uploaded_file: Uploaded file object

    Returns:
        dict: File information
    """
    return {
        'name': uploaded_file.name,
        'size': len(uploaded_file.getvalue()),
        'type': uploaded_file.type
    }

